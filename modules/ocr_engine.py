"""
OCR Engine — extracts raw text from PDFs (native + scanned) and images.
"""

import io
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text(file_path: str) -> dict:
    """
    Master text extractor. Dispatches to the correct engine
    based on file extension.

    Returns:
        {
            "raw_text": str,
            "method": "pdfplumber" | "ocr" | "docx",
            "page_count": int,
            "ocr_used": bool,
        }
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext in (".doc", ".docx"):
        return _extract_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        return _extract_image_ocr(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_path: str) -> dict:
    """Try native text extraction first; fall back to OCR for scanned pages."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not available, falling back to OCR")
        return _pdf_ocr_fallback(file_path)

    full_text = []
    page_count = 0
    ocr_used = False

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text()
            if text and len(text.strip()) > 20:
                full_text.append(text)
            else:
                # Scanned page — use OCR
                ocr_text = _ocr_page_image(page)
                if ocr_text:
                    full_text.append(ocr_text)
                    ocr_used = True

    raw = "\n\n".join(full_text).strip()
    method = "ocr" if ocr_used and not any(
        p for p in full_text if len(p) > 100
    ) else "pdfplumber"

    return {
        "raw_text": raw,
        "method": method,
        "page_count": page_count,
        "ocr_used": ocr_used,
    }


def _ocr_page_image(page) -> str:
    """Convert a pdfplumber page to image and OCR it."""
    try:
        import pytesseract
        from PIL import Image

        page_image = page.to_image(resolution=300)
        pil_img = page_image.original
        text = pytesseract.image_to_string(pil_img, lang="eng")
        return text.strip()
    except Exception as e:
        logger.error(f"OCR on PDF page failed: {e}")
        return ""


def _pdf_ocr_fallback(file_path: str) -> dict:
    """Full OCR pipeline using pdf2image when pdfplumber unavailable."""
    try:
        import pytesseract
        from pdf2image import convert_from_path

        pages = convert_from_path(file_path, dpi=300)
        texts = [pytesseract.image_to_string(p, lang="eng") for p in pages]
        return {
            "raw_text": "\n\n".join(texts).strip(),
            "method": "ocr",
            "page_count": len(pages),
            "ocr_used": True,
        }
    except Exception as e:
        logger.error(f"PDF OCR fallback failed: {e}")
        return {"raw_text": "", "method": "error", "page_count": 0, "ocr_used": False}


def _extract_docx(file_path: str) -> dict:
    """Extract text from .doc/.docx using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        return {
            "raw_text": "\n".join(paragraphs).strip(),
            "method": "docx",
            "page_count": 1,
            "ocr_used": False,
        }
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return {"raw_text": "", "method": "error", "page_count": 0, "ocr_used": False}


def _extract_image_ocr(file_path: str) -> dict:
    """Extract text from image files using Tesseract."""
    try:
        import pytesseract
        from PIL import Image, ImageFilter, ImageEnhance

        img = Image.open(file_path)

        # Pre-processing for better OCR accuracy
        img = img.convert("L")                        # Grayscale
        img = img.filter(ImageFilter.SHARPEN)         # Sharpen
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)                   # Boost contrast

        text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
        return {
            "raw_text": text.strip(),
            "method": "ocr",
            "page_count": 1,
            "ocr_used": True,
        }
    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        return {"raw_text": "", "method": "error", "page_count": 0, "ocr_used": False}
